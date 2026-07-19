import streamlit as st
import os
import re
import pandas as pd
import openpyxl
import docx
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from rank_bm25 import BM25Okapi
from groq import Groq

# ----------------------------------------------------
# 1. Page & Layout Configuration
# ----------------------------------------------------
st.set_page_config(
    page_title="NDSP Domain-Expert Assistant",
    page_icon="🥛",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("🥛 NDSP Domain-Expert Assistant")
st.caption("Ground your queries directly inside Breed Improvement, Animal Nutrition, Bulking, and EFA parameters.")

# Initialize session states for storing processed application data
if "all_chunks" not in st.session_state:
    st.session_state.all_chunks = None
if "sheet_frames" not in st.session_state:
    st.session_state.sheet_frames = None
if "faiss_index" not in st.session_state:
    st.session_state.faiss_index = None
if "bm25" not in st.session_state:
    st.session_state.bm25 = None
if "messages" not in st.session_state:
    st.session_state.messages = []

# ----------------------------------------------------
# 2. Cached Models & Resource Allocations
# ----------------------------------------------------
@st.cache_resource(show_spinner="Loading Embedding Architecture (BAAI/bge-small-en-v1.5)...")
def load_embedding_model():
    return SentenceTransformer("BAAI/bge-small-en-v1.5")

embed_model = load_embedding_model()

# ----------------------------------------------------
# 3. Document Ingestion Extractors (In-Memory Processing)
# ----------------------------------------------------
def ingest_pdf(uploaded_file, paras_per_block=6):
    file_name = uploaded_file.name
    reader = PdfReader(uploaded_file)
    chunks = []
    buf = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            lines = text.split('\n')
            for line in lines:
                stripped_line = line.strip()
                if stripped_line:
                    buf.append(stripped_line)
                    if len(buf) >= paras_per_block:
                        chunks.append({
                            "text": f"[SOURCE: {file_name} | PAGE {page_num + 1} | SECTION TEXT]\n" + "\n".join(buf),
                            "file": file_name,
                            "sheet": f"page_{page_num + 1}",
                            "row_start": None
                        })
                        buf = []
    if buf:
        chunks.append({
            "text": f"[SOURCE: {file_name} | PAGE {len(reader.pages)} | SECTION TEXT]\n" + "\n".join(buf),
            "file": file_name,
            "sheet": f"page_{len(reader.pages)}",
            "row_start": None
        })
    return chunks

def load_sheet_df(wb, sheet):
    ws = wb[sheet]
    data = list(ws.iter_rows(values_only=True))
    df = pd.DataFrame(data)
    df = df.dropna(axis=0, how="all").dropna(axis=1, how="all")
    return df

def fmt_val(v):
    if isinstance(v, float):
        return f"{v:.4g}"
    return str(v)

def df_to_text_blocks(df, file_name, sheet, rows_per_block=35):
    blocks = []
    n = len(df)
    for start in range(0, n, rows_per_block):
        sub = df.iloc[start:start + rows_per_block]
        lines = []
        for _, row in sub.iterrows():
            vals = [fmt_val(v) for v in row.tolist() if pd.notna(v)]
            if vals:
                lines.append(" | ".join(vals))
        if lines:
            header = f"[SOURCE: {file_name} | SHEET: {sheet} | rows {start + 1} - {min(start + rows_per_block, n)}]"
            blocks.append({
                "text": header + "\n" + "\n".join(lines),
                "file": file_name,
                "sheet": sheet,
                "row_start": start + 1
            })
    return blocks

def ingest_excel(uploaded_file):
    file_name = uploaded_file.name
    wb = openpyxl.load_workbook(uploaded_file, data_only=True)
    chunks = []
    sheet_frames = {}
    for sheet in wb.sheetnames:
        try:
            df = load_sheet_df(wb, sheet)
        except Exception as e:
            continue
        if df.empty:
            continue
        sheet_frames[(file_name, sheet)] = df
        chunks.extend(df_to_text_blocks(df, file_name, sheet))
    return chunks, sheet_frames

def ingest_docx(uploaded_file, paras_per_block=6):
    file_name = uploaded_file.name
    d = docx.Document(uploaded_file)
    chunks = []
    buf = []
    for p in d.paragraphs:
        if p.text.strip():
            buf.append(p.text.strip())
            if len(buf) >= paras_per_block:
                chunks.append({
                    "text": f"[SOURCE: {file_name} | SECTION TEXT]\n" + "\n".join(buf),
                    "file": file_name,
                    "sheet": None,
                    "row_start": None
                })
                buf = []
    if buf:
        chunks.append({
            "text": f"[SOURCE: {file_name} | SECTION TEXT]\n" + "\n".join(buf),
            "file": file_name,
            "sheet": None,
            "row_start": None
        })
    for ti, table in enumerate(d.tables):
        rows_txt = []
        for row in table.rows:
            cells_txt = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells_txt:
                rows_txt.append(" | ".join(cells_txt))
        if rows_txt:
            chunks.append({
                "text": f"[SOURCE: {file_name} | TABLE {ti + 1}]\n" + "\n".join(rows_txt),
                "file": file_name,
                "sheet": f"table_{ti + 1}",
                "row_start": None
            })
    return chunks

# ----------------------------------------------------
# 4. Search and Retrieval Frameworks
# ----------------------------------------------------
def semantic_search(query, index, k=8):
    q_emb = embed_model.encode([query], normalize_embeddings=True)
    scores, idxs = index.search(np.array(q_emb, dtype="float32"), k)
    return list(idxs[0])

def bm25_search(query, bm25, k=8):
    tok = re.findall(r"\w+", query.lower())
    scores = bm25.get_scores(tok)
    top = np.argsort(scores)[::-1][:k]
    return list(top)

def hybrid_retrieve(query, all_chunks, index, bm25, k=6, pool=15):
    sem = semantic_search(query, index, pool)
    kw = bm25_search(query, bm25, pool)
    rank_score = {}
    for rank, idx in enumerate(sem):
        rank_score[idx] = rank_score.get(idx, 0) + 1.0 / (rank + 1)
    for rank, idx in enumerate(kw):
        rank_score[idx] = rank_score.get(idx, 0) + 1.0 / (rank + 1)
    ranked = sorted(rank_score.items(), key=lambda x: -x[1])[:k]
    return [all_chunks[i] for i, _ in ranked]

def structured_lookup(query, sheet_frames, max_hits=5):
    terms = [t for t in re.findall(r"\w+", query.lower()) if len(t) > 3]
    if not terms:
        return []
    hits = []
    for (file, sheet), df in sheet_frames.items():
        for ridx, row in df.iterrows():
            row_vals = [str(v) for v in row.tolist() if pd.notna(v)]
            row_str = " ".join(row_vals).lower()
            if any(t in row_str for t in terms):
                hits.append(f"[SOURCE: {file} | SHEET: {sheet} | row {ridx + 1}] " + " | ".join(row_vals))
            if len(hits) >= max_hits:
                break
        if len(hits) >= max_hits:
            break
    return hits

# ----------------------------------------------------
# 5. Application Setup Sidebar Inputs
# ----------------------------------------------------
with st.sidebar:
    st.header("🔑 Authentication Setup")
    groq_api_key = st.text_input("Enter Groq API Key", type="password", help="Grab an API key via console.groq.com")
    
    st.header("📁 Document Ingestion")
    uploaded_files = st.file_uploader(
        "Upload NDSP Project Files",
        type=["pdf", "xlsx", "docx"],
        accept_multiple_files=True
    )
    
    process_btn = st.button("Build Knowledge Base", use_container_width=True)

# ----------------------------------------------------
# 6. File Processing Pipeline Execution
# ----------------------------------------------------
if process_btn:
    if not groq_api_key:
        st.sidebar.error("Please supply a valid Groq API Key first.")
    elif not uploaded_files:
        st.sidebar.error("Please select one or more files to analyze.")
    else:
        # Utilize container tracking status framework to display progress clearly
        with st.status("Constructing Knowledge Engine...", expanded=True) as status:
            local_chunks = []
            local_sheets = {}
            
            for f in uploaded_files:
                status.write(f"🔄 Parsing elements inside: **{f.name}**")
                if f.name.lower().endswith('.pdf'):
                    local_chunks.extend(ingest_pdf(f))
                elif f.name.lower().endswith('.xlsx'):
                    c, sf = ingest_excel(f)
                    local_chunks.extend(c)
                    local_sheets.update(sf)
                elif f.name.lower().endswith('.docx'):
                    local_chunks.extend(ingest_docx(f))
            
            if not local_chunks:
                status.update(label="❌ Failed processing text structures.", state="error")
            else:
                status.write(f"⚙️ Vectorizing {len(local_chunks)} split paragraphs/table segments...")
                texts = [c["text"] for c in local_chunks]
                embeddings = embed_model.encode(texts, batch_size=32, show_progress_bar=False, normalize_embeddings=True)
                
                dim = embeddings.shape[1]
                index = faiss.IndexFlatIP(dim)
                index.add(np.array(embeddings, dtype="float32"))
                
                status.write("🛠️ Structuring inverted Keyword mapping (BM25)...")
                tokenized = [re.findall(r"\w+", t.lower()) for t in texts]
                bm25 = BM25Okapi(tokenized)
                
                # Save processed items to global instance memory safely
                st.session_state.all_chunks = local_chunks
                st.session_state.sheet_frames = local_sheets
                st.session_state.faiss_index = index
                st.session_state.bm25 = bm25
                
                status.update(label="🚀 Core Knowledge Engine Ready!", state="complete")
                st.sidebar.success(f"Successfully processed {len(uploaded_files)} files!")

# ----------------------------------------------------
# 7. Main Interactive Chat Interface Container
# ----------------------------------------------------
if st.session_state.faiss_index is not None:
    # Clear history options utility metric
    if st.button("🧹 Clear Chat History"):
        st.session_state.messages = []
        st.rerun()

    # Draw historic loops safely
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            
    # Capture queries sent inside runtime interfaces
    if user_query := st.chat_input("Ask a question about your uploaded data..."):
        with st.chat_message("user"):
            st.markdown(user_query)
        st.session_state.messages.append({"role": "user", "content": user_query})
        
        if not groq_api_key:
            st.error("Missing API Key! Please enter it in the left sidebar configuration window.")
        else:
            with st.chat_message("assistant"):
                with st.spinner("Retrieving target resources & evaluating context logic..."):
                    try:
                        # Extract components from global state safely
                        chunks = st.session_state.all_chunks
                        idx = st.session_state.faiss_index
                        bm25_obj = st.session_state.bm25
                        sheets = st.session_state.sheet_frames
                        
                        # Apply hybrid search algorithms
                        retrieved = hybrid_retrieve(user_query, chunks, idx, bm25_obj, k=6)
                        struct_hits = structured_lookup(user_query, sheets)
                        
                        context_parts = [c["text"] for c in retrieved]
                        if struct_hits:
                            context_parts.append("[DIRECT SHEET MATCHES]\n" + "\n".join(struct_hits))
                        
                        full_context = "\n\n---\n\n".join(context_parts)
                        
                        # Establish interaction connections through Groq endpoints
                        client = Groq(api_key=groq_api_key)
                        messages_payload = [
                            {"role": "system", "content": "You are a Domain-Expert Assistant. Answer the question comprehensively based strictly on the provided context sources. Always ground names and metrics transparently."},
                            {"role": "user", "content": f"CONTEXT:\n{full_context}\n\nQUESTION: {user_query}"}
                        ]
                        
                        completion = client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=messages_payload,
                            temperature=0.1,
                            max_tokens=1000
                        )
                        
                        response_text = completion.choices[0].message.content
                        st.markdown(response_text)
                        st.session_state.messages.append({"role": "assistant", "content": response_text})
                        
                    except Exception as err:
                        st.error(f"Runtime execution failure: {err}")
else:
    # Onboard instruction alerts when active structures don't exist yet
    st.info("👈 Set your Groq API Key and upload your NDSP research documentation (.pdf, .xlsx, .docx) in the sidebar to begin.")
